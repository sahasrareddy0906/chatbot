import io
from pypdf import PdfReader
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str | None:
    """
    Extract all text from a PDF file's raw bytes.
    """
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        full_text = "\n".join(text_parts)
        return full_text.strip() if full_text.strip() else None

    except Exception as e:
        print(f"❌ PDF extraction failed: {e}")
        return None


def extract_text_from_docx(file_bytes: bytes) -> str | None:
    """
    Extract all text from a DOCX file's raw bytes.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables (some resumes use table layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        full_text = "\n".join(text_parts)
        return full_text.strip() if full_text.strip() else None

    except Exception as e:
        print(f"❌ DOCX extraction failed: {e}")
        return None


def extract_resume_text(file_bytes: bytes,
                        filename: str) -> dict:
    """
    Detect file type and extract text accordingly.
    Returns dict with text and any warnings.
    """
    ext = filename.split(".")[-1].lower()

    if ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext in ["doc", "docx"]:
        text = extract_text_from_docx(file_bytes)
    else:
        return {
            "success": False,
            "error":   f"Unsupported file type: {ext}"
        }

    if not text:
        return {
            "success": False,
            "error":   "Could not extract text. "
                      "File may be scanned/image-based or corrupted."
        }

    # Basic sanity check — resume should have reasonable length
    if len(text) < 100:
        return {
            "success": False,
            "error":   "Extracted text too short. "
                      "File may not contain readable text."
        }

    return {
        "success": True,
        "text":    text,
        "length":  len(text)
    }